"""
exporter.py

Exports markdown blog articles to clean, professional Microsoft Word (.docx)
and Adobe PDF (.pdf) documents.

Excludes evaluator scores, metrics, radar charts, and pipeline logs.
Includes styled headings, paragraphs, formatted lists, tables, callouts, and embedded images.
"""

from __future__ import annotations

import io
import os
import re
from pathlib import Path
from typing import Optional
from PIL import Image as PILImage

# ── Clean / Sanitization Helper ───────────────────────────────────────────────

def clean_blog_markdown(markdown: str) -> str:
    """
    Strips internal evaluation scores, baseline comparison tables,
    JSON score blocks, and pipeline metadata from the markdown text.
    """
    if not markdown:
        return ""

    text = markdown.strip()

    # Remove JSON score blocks (e.g. ```json ... "language_score": ... ```)
    text = re.sub(r"```json\s*\{[\s\S]*?\}\s*```", "", text, flags=re.IGNORECASE)

    # Remove score table markers or evaluation summaries if present
    patterns = [
        r"###?\s*Evaluator Scores[\s\S]*?(?=\n##|\Z)",
        r"###?\s*Baseline Scores[\s\S]*?(?=\n##|\Z)",
        r"###?\s*Evaluation Summary[\s\S]*?(?=\n##|\Z)",
        r"<table[\s\S]*?class=['\"]sct['\"][\s\S]*?</table>",
    ]
    for pattern in patterns:
        text = re.sub(pattern, "", text, flags=re.IGNORECASE)

    # Clean redundant blank lines
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _resolve_image_path(img_ref: str) -> Optional[Path]:
    """Finds an image file on the local filesystem from markdown image target."""
    clean_ref = img_ref.strip().strip("'").strip('"')
    p = Path(clean_ref)
    if p.exists() and p.is_file():
        return p

    # Check relative to workspace root or generated_images
    candidates = [
        Path.cwd() / clean_ref,
        Path.cwd() / "generated_images" / clean_ref,
        Path.cwd() / "generated_images" / p.name,
        Path.cwd() / "data" / "custom_images" / p.name,
    ]
    for cand in candidates:
        if cand.exists() and cand.is_file():
            return cand
    return None


# ── DOCX Export ───────────────────────────────────────────────────────────────

def export_to_docx(markdown_text: str, title: str = "Blog Article", default_title: str = "Blog Article") -> bytes:
    """
    Converts markdown blog text into a styled Microsoft Word (.docx) document.
    Returns bytes suitable for downloading via Streamlit.
    """
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc_title = title or default_title or "Blog Article"
    cleaned = clean_blog_markdown(markdown_text)
    doc = Document()

    # Set standard margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Palette
    COLOR_PRIMARY = RGBColor(15, 23, 42)    # Slate 900
    COLOR_SECONDARY = RGBColor(30, 41, 59)  # Slate 800
    COLOR_MUTED = RGBColor(100, 116, 139)   # Slate 500
    COLOR_BODY = RGBColor(30, 41, 59)

    # Configure Normal Style
    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = COLOR_BODY

    lines = cleaned.splitlines()
    i = 0
    in_code_block = False
    code_lines = []

    def _add_styled_inline(paragraph, text: str, italic: bool = False):
        """Parse basic bold/italic inline markdown in paragraph."""
        tokens = re.split(r"(\*\*.*?\*\*|\*.*?\*|`.*?`)", text)
        for tok in tokens:
            if not tok:
                continue
            if tok.startswith("**") and tok.endswith("**") and len(tok) >= 4:
                run = paragraph.add_run(tok[2:-2])
                run.bold = True
                if italic:
                    run.italic = True
            elif tok.startswith("*") and tok.endswith("*") and len(tok) >= 2:
                run = paragraph.add_run(tok[1:-1])
                run.italic = True
            elif tok.startswith("`") and tok.endswith("`") and len(tok) >= 2:
                run = paragraph.add_run(tok[1:-1])
                run.font.name = "Consolas"
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(180, 83, 9)
            else:
                run = paragraph.add_run(tok)
                if italic:
                    run.italic = True

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Handle Code blocks
        if stripped.startswith("```"):
            if in_code_block:
                code_text = "\n".join(code_lines)
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.left_indent = Inches(0.3)
                run = p.add_run(code_text)
                run.font.name = "Consolas"
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(51, 65, 85)
                in_code_block = False
                code_lines = []
            else:
                in_code_block = True
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Skip empty lines
        if not stripped:
            i += 1
            continue

        # Image detection: ![alt](url)
        img_match = re.search(r"^!\[(?P<alt>[^\]]*)\]\((?P<url>[^)]+)\)", stripped)
        if img_match:
            alt_text = img_match.group("alt")
            img_target = img_match.group("url")
            img_path = _resolve_image_path(img_target)

            if img_path:
                try:
                    p_img = doc.add_paragraph()
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_img.paragraph_format.space_before = Pt(8)
                    p_img.paragraph_format.space_after = Pt(2)
                    run_img = p_img.add_run()
                    run_img.add_picture(str(img_path), width=Inches(5.2))
                except Exception:
                    pass

            caption_text = alt_text
            if i + 1 < len(lines) and lines[i+1].strip().startswith("*") and lines[i+1].strip().endswith("*"):
                i += 1
                caption_text = lines[i].strip().strip("*").strip()

            if caption_text:
                p_cap = doc.add_paragraph()
                p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_cap.paragraph_format.space_before = Pt(0)
                p_cap.paragraph_format.space_after = Pt(8)
                run_cap = p_cap.add_run(caption_text)
                run_cap.italic = True
                run_cap.font.size = Pt(9.5)
                run_cap.font.color.rgb = COLOR_MUTED

            i += 1
            continue

        # Headings
        if stripped.startswith("# "):
            h = doc.add_heading(level=1)
            h.paragraph_format.space_before = Pt(16)
            h.paragraph_format.space_after = Pt(6)
            h.paragraph_format.keep_with_next = True
            run = h.add_run(stripped[2:].strip())
            run.font.name = "Calibri"
            run.font.size = Pt(20)
            run.font.bold = True
            run.font.color.rgb = COLOR_PRIMARY
            i += 1
            continue

        if stripped.startswith("## "):
            h = doc.add_heading(level=2)
            h.paragraph_format.space_before = Pt(14)
            h.paragraph_format.space_after = Pt(4)
            h.paragraph_format.keep_with_next = True
            run = h.add_run(stripped[3:].strip())
            run.font.name = "Calibri"
            run.font.size = Pt(15)
            run.font.bold = True
            run.font.color.rgb = COLOR_SECONDARY
            i += 1
            continue

        if stripped.startswith("### "):
            h = doc.add_heading(level=3)
            h.paragraph_format.space_before = Pt(10)
            h.paragraph_format.space_after = Pt(3)
            h.paragraph_format.keep_with_next = True
            run = h.add_run(stripped[4:].strip())
            run.font.name = "Calibri"
            run.font.size = Pt(13)
            run.font.bold = True
            run.font.color.rgb = COLOR_SECONDARY
            i += 1
            continue

        # Bullet List Items
        if stripped.startswith(("- ", "* ", "+ ")):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(3)
            _add_styled_inline(p, stripped[2:].strip())
            i += 1
            continue

        # Numbered List Items
        num_match = re.match(r"^(\d+)\.\s+(.*)", stripped)
        if num_match:
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(3)
            _add_styled_inline(p, num_match.group(2).strip())
            i += 1
            continue

        # Blockquote
        if stripped.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(6)
            _add_styled_inline(p, stripped[1:].strip(), italic=True)
            i += 1
            continue

        # Regular Paragraph
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        _add_styled_inline(p, stripped)
        i += 1

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ── PDF Export ────────────────────────────────────────────────────────────────

def export_to_pdf(markdown_text: str, title: str = "Blog Article", default_title: str = "Blog Article") -> bytes:
    """
    Converts markdown blog text into a publication-quality PDF document.
    Returns bytes suitable for downloading via Streamlit.
    """
    from reportlab.lib.pagesizes import letter
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Image as RLImage, KeepTogether
    )
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.pdfgen import canvas

    doc_title = title or default_title or "Blog Article"
    cleaned = clean_blog_markdown(markdown_text)

    class NumberedCanvas(canvas.Canvas):
        def __init__(self, *args, **kwargs):
            super().__init__(*args, **kwargs)
            self._saved_page_states = []

        def showPage(self):
            self._saved_page_states.append(dict(self.__dict__))
            self._startPage()

        def save(self):
            num_pages = len(self._saved_page_states)
            for state in self._saved_page_states:
                self.__dict__.update(state)
                self.draw_page_number(num_pages)
                super().showPage()
            super().save()

        def draw_page_number(self, page_count):
            self.saveState()
            self.setFont("Helvetica", 8.5)
            self.setFillColor(colors.HexColor("#64748B"))

            # Header rule & title
            self.setLineWidth(0.5)
            self.setStrokeColor(colors.HexColor("#E2E8F0"))
            self.line(54, 792 - 40, 612 - 54, 792 - 40)
            self.drawString(54, 792 - 34, default_title[:60])

            # Footer
            self.line(54, 45, 612 - 54, 45)
            page_text = f"Page {self._pageNumber} of {page_count}"
            self.drawRightString(612 - 54, 32, page_text)
            self.restoreState()

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        leftMargin=54,
        rightMargin=54,
        topMargin=54,
        bottomMargin=54,
    )

    styles = getSampleStyleSheet()

    style_title = ParagraphStyle(
        "BlogTitle",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=20,
        leading=24,
        textColor=colors.HexColor("#0F172A"),
        spaceBefore=10,
        spaceAfter=10,
        alignment=TA_LEFT,
    )
    style_h2 = ParagraphStyle(
        "BlogH2",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=14,
        leading=18,
        textColor=colors.HexColor("#1E293B"),
        spaceBefore=12,
        spaceAfter=5,
        keepWithNext=True,
    )
    style_h3 = ParagraphStyle(
        "BlogH3",
        parent=styles["Heading3"],
        fontName="Helvetica-Bold",
        fontSize=11.5,
        leading=15,
        textColor=colors.HexColor("#334155"),
        spaceBefore=9,
        spaceAfter=3,
        keepWithNext=True,
    )
    style_body = ParagraphStyle(
        "BlogBody",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=10,
        leading=14.5,
        textColor=colors.HexColor("#1E293B"),
        spaceBefore=0,
        spaceAfter=6,
        alignment=TA_LEFT,
    )
    style_bullet = ParagraphStyle(
        "BlogBullet",
        parent=style_body,
        leftIndent=15,
        firstLineIndent=-10,
        spaceAfter=3,
    )
    style_caption = ParagraphStyle(
        "BlogCaption",
        parent=styles["Italic"],
        fontName="Helvetica-Oblique",
        fontSize=8.5,
        leading=11,
        textColor=colors.HexColor("#64748B"),
        alignment=TA_CENTER,
        spaceBefore=2,
        spaceAfter=8,
    )
    style_blockquote = ParagraphStyle(
        "BlogQuote",
        parent=style_body,
        fontName="Helvetica-Oblique",
        leftIndent=20,
        rightIndent=20,
        textColor=colors.HexColor("#475569"),
        spaceBefore=4,
        spaceAfter=6,
    )

    def _md_to_rl_xml(text: str) -> str:
        """Convert markdown bold/italic/code syntax to ReportLab XML tags."""
        t = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        t = re.sub(r"\*\*(.*?)\*\*", r"<b>\1</b>", t)
        t = re.sub(r"\*(.*?)\*", r"<i>\1</i>", t)
        t = re.sub(r"`(.*?)`", r"<font face='Courier' color='#B45309'>\1</font>", t)
        return t

    story = []
    lines = cleaned.splitlines()
    i = 0
    in_code_block = False
    code_lines = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Handle Code blocks
        if stripped.startswith("```"):
            if in_code_block:
                code_text = "\n".join(code_lines)
                xml_code = _md_to_rl_xml(code_text)
                style_code = ParagraphStyle(
                    "BlogCode",
                    parent=styles["Normal"],
                    fontName="Courier",
                    fontSize=8.5,
                    leading=11.5,
                    textColor=colors.HexColor("#334155"),
                    backColor=colors.HexColor("#F8FAFC"),
                    borderPadding=6,
                    spaceBefore=4,
                    spaceAfter=6,
                )
                story.append(Paragraph(xml_code.replace("\n", "<br/>"), style_code))
                in_code_block = False
                code_lines = []
            else:
                in_code_block = True
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        # Image detection
        img_match = re.search(r"^!\[(?P<alt>[^\]]*)\]\((?P<url>[^)]+)\)", stripped)
        if img_match:
            alt_text = img_match.group("alt")
            img_target = img_match.group("url")
            img_path = _resolve_image_path(img_target)

            caption_text = alt_text
            if i + 1 < len(lines) and lines[i+1].strip().startswith("*") and lines[i+1].strip().endswith("*"):
                i += 1
                caption_text = lines[i].strip().strip("*").strip()

            if img_path:
                try:
                    with PILImage.open(str(img_path)) as pil_img:
                        orig_w, orig_h = pil_img.size

                    max_w = 460
                    max_h = 260
                    scale = min(max_w / orig_w, max_h / orig_h, 1.0)
                    target_w = orig_w * scale
                    target_h = orig_h * scale

                    img_flow = RLImage(str(img_path), width=target_w, height=target_h)
                    img_flow.hAlign = "CENTER"

                    img_elements = [Spacer(1, 4), img_flow]
                    if caption_text:
                        img_elements.append(Paragraph(_md_to_rl_xml(caption_text), style_caption))
                    story.append(KeepTogether(img_elements))
                except Exception:
                    pass
            elif caption_text:
                story.append(Paragraph(_md_to_rl_xml(f"[{caption_text}]"), style_caption))

            i += 1
            continue

        # Headings
        if stripped.startswith("# "):
            story.append(Paragraph(_md_to_rl_xml(stripped[2:].strip()), style_title))
            i += 1
            continue

        if stripped.startswith("## "):
            story.append(Paragraph(_md_to_rl_xml(stripped[3:].strip()), style_h2))
            i += 1
            continue

        if stripped.startswith("### "):
            story.append(Paragraph(_md_to_rl_xml(stripped[4:].strip()), style_h3))
            i += 1
            continue

        # Bullet List Items
        if stripped.startswith(("- ", "* ", "+ ")):
            item_text = f"&bull;&nbsp; {_md_to_rl_xml(stripped[2:].strip())}"
            story.append(Paragraph(item_text, style_bullet))
            i += 1
            continue

        # Numbered List Items
        num_match = re.match(r"^(\d+)\.\s+(.*)", stripped)
        if num_match:
            item_text = f"{num_match.group(1)}.&nbsp; {_md_to_rl_xml(num_match.group(2).strip())}"
            story.append(Paragraph(item_text, style_bullet))
            i += 1
            continue

        # Blockquote
        if stripped.startswith(">"):
            story.append(Paragraph(_md_to_rl_xml(stripped[1:].strip()), style_blockquote))
            i += 1
            continue

        # Normal Paragraph
        story.append(Paragraph(_md_to_rl_xml(stripped), style_body))
        i += 1

    doc.build(story, canvasmaker=NumberedCanvas)
    buffer.seek(0)
    return buffer.getvalue()
